# -*- coding: utf-8 -*-
import openpyxl
import codecs
from openpyxl.utils import get_column_letter
import os
import sys
import xlwt
import xlrd
import time
import datetime
import math
import random
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
 


#第一部分
#######################################################
#以下产生从txt产生excel表格                           
#######################################################
def readInAFile(f):
    '''读入一个txt文件, 转成df, 用于之后合并成一个大的df, 这个大的df就是最后要呈现的结果'''
    tmp = pd.read_csv(f,sep='\t', header=None)
    print('\n')
    #print(tmp)
    return tmp



def convert_txt_to_df(apath):
    '''这个函数就是合并所有txt得到最后df的函数'''

    res = None
    for f in os.listdir(apath):
        if f.endswith('.txt') is False: continue
        #print(f)
        afile = os.path.join(apath,f)
        tmpres = readInAFile(afile)
        #print(tmpres)

        if res is None: res = tmpres
        #else: res = res.append(tmpres)
        else: res = pd.concat([res, tmpres])

    return res



def my_str_cat(a, b):
     '''拼接y和m'''

     if len(str(b)) == 1:
         b = '0'+str(b)
     return str(a) + '-' + str(b)



def com_sum(adf): 
    '''计算总计的函数'''

    #com sum
    adf['sum'] = adf.apply(lambda x: x.sum(), axis=1 )



def com_ratio(adf):
    '''计算占比的函数'''
    #com ratio
    #put the ratio to the right of the scale block
    #pass

    colsName = adf.columns
    #print(colsName)

    for col in colsName[:-1]:
        adf[col+'_ratio'] = adf[col]/adf['sum']

    adf.rename(columns={'sum':'总计'}, inplace = True)
    print(adf.iloc[:,adf.shape[1]//2:])



def getYM():
    '''从当前时间获取月报月份'''

    now = datetime.datetime.now()
    year = now.year
    month = now.month
    month = month-1
    print(year, month)
    thisM = month
    lastM = thisM-1
    lastY = year - 1

    if len(str(thisM))==1: 
        thisM = '0'+str(thisM)
    if len(str(lastM))==1: 
        lastM = '0'+str(lastM)

    thisYthisM = str(year)+'-'+str(thisM)
    thisYlastM = str(year)+'-'+str(lastM)
    
    lastYthisM = str(lastY)+'-'+str(thisM)
    lastYlastM = str(lastY)+'-'+str(lastM)

    llYthisM = str(year-2)+'-'+str(thisM)

    print(thisYthisM, thisYlastM, lastYthisM, lastYlastM, llYthisM)

    return thisYthisM, thisYlastM, lastYthisM, lastYlastM, llYthisM



def com_rate(adf,ym=None):
    '''com rate
    put the rate to the upper of the scale and ratio blocks
    '''

    tYtM, tYlM, lYtM, lYlM, llYtM = getYM()

    adf.loc['环比增幅'] = adf.apply(lambda x: round(x.loc[tYtM]/x.loc[tYlM]-1,4))
   
    adf.loc['环比增量'] = adf.apply(lambda x: x.loc[tYtM]-x.loc[tYlM])
    adf.loc['ph1'] = ['' for i in range(adf.shape[1])]

    adf.loc['同期环比增幅'] = adf.apply(lambda x: round(x.loc[lYtM]/x.loc[lYlM]-1,4))
    adf.loc['同期环比增量'] = adf.apply(lambda x: x.loc[lYtM]-x.loc[lYlM])
    adf.loc['ph2'] = ['' for i in range(adf.shape[1])]

    adf.loc['同比增幅'] = adf.apply(lambda x: round(x.loc[tYtM]/x.loc[lYtM]-1,4))
    adf.loc['同比增量'] = adf.apply(lambda x: x.loc[tYtM]-x.loc[lYtM])
    adf.loc['ph3'] = ['' for i in range(adf.shape[1])]
    adf.loc['同期同比增幅'] = adf.apply(lambda x: round(x.loc[lYtM]/x.loc[llYtM]-1,4))
    adf.loc['同期同比增量'] = adf.apply(lambda x: x.loc[lYtM]-x.loc[llYtM])

    foreward = adf[:-11]
    backward = adf[-11:]

    #header = ['N_'+str(i) for i in range(adf.shape[1])]
    #header = np.array(header)
    #print(header.shape)
    #header = header.reshape(1,23)

    #colNames = [col for col in adf.columns]
    #header = pd.DataFrame(header, columns = colNames) 
   
    #adf = pd.concat([backward,header,foreward])
    adf = pd.concat([backward,foreward])
    return adf



def my_round(adf):
    for i in [0,3,6,9]:
        adf.iloc[i,:] = adf.iloc[i,:].apply(lambda x: format(x, '.2%')) 
        #for j in range(adf.shape[1]//2+1,adf.shape[1]):
        #    adf.iloc[i,j] = round(float(adf.iloc[i,j]),8) 
    for i in [1,4,7,10]:
        adf.iloc[i,adf.shape[1]//2+1:] = adf.iloc[i,adf.shape[1]//2+1:].apply(lambda x: format(x, '.2%'))

    print(adf)
    return adf



def sortDF(adf):
    '''上一个函数得到的结果并没有按照时间来排序'''
    #pass
    adfcols = adf.shape[1]
    adfrows = adf.shape[0]
    adf.columns = ['col_'+str(i) for i in range(adfcols)]
    adf.index = ['row_'+str(i) for i in range(adfrows)]
    adf['col_0'] = pd.to_datetime(adf['col_0'])
    adf['col_0'] = pd.to_datetime(adf['col_0'],format='%Y%m')
    adf = adf.sort_values(by='col_0',ascending=False)
   
    # convert %Y-%m-%m to %Y-%m
    adf['y'] = adf['col_0'].dt.year
    adf['m'] = adf['col_0'].dt.month
    adf['ym'] = adf.apply(lambda adf: my_str_cat(adf['y'], adf['m']), axis=1)
    adf.drop(['col_0'],axis=1,inplace=True)
    adf.drop(['y'],axis=1,inplace=True)
    adf.drop(['m'],axis=1,inplace=True)

    adf.rename(columns={'ym':'时间'}, inplace = True)

    #adf.rename(columns={'col_0':'时间'}, inplace = True)
    adf.set_index(['时间'], inplace = True)
    
    com_sum(adf)
    com_ratio(adf)
    adf = com_rate(adf)
    #adf = my_round(adf)

    return adf



def outputExcel(excelName, sheetName, adf):
    '''通过df生成一个excel'''
    adf.to_excel(excelName,sheet_name=sheetName)
 


########################################################
#以下是修饰excel表所用，比如居中等                     
########################################################
def decorate():
    #居中，有时候需要注意index为左对齐，但是大部分为居中对齐，竖直方向也为居中对齐
    #调整字体大小，首行为12，其他为11，当月和上月为blod，同期当月和上月为bold，同期的同期当月为bold
    #按照列 设置规则
    #pass
    font = xlwt.Font()
    font.name = '微软雅黑'
    font.colour_index = 0 #black

    font_header = xlwt.Font()
    font_header.name = '微软雅黑'
    font_header.colour_index = 2 #red
    font_header.bold= True

    font_index = xlwt.Font()
    font_index.name = '微软雅黑'
    font_index.colour_index = 0 #red
    #font_index.bold= True

    font_col = xlwt.Font()
    font_col.name = '微软雅黑'
    font_col.colour_index = 1 # white
    font_col.bold= True

    alignment = xlwt.Alignment()
    alignment.horz = xlwt.Alignment.HORZ_CENTER
    alignment.vert = xlwt.Alignment.VERT_CENTER

    #上框线
    borders = xlwt.Borders()
    borders.top = xlwt.Borders.THIN

    #下框线
    borders2 = xlwt.Borders()
    borders2.bottom = xlwt.Borders.THIN

    #右框线
    borders3 = xlwt.Borders()
    borders3.right = xlwt.Borders.THIN

    #右下框线
    borders4 = xlwt.Borders()
    borders4.right = xlwt.Borders.THIN
    borders4.bottom = xlwt.Borders.THIN

    #右上框线
    borders5 = xlwt.Borders()
    borders5.right = xlwt.Borders.THIN
    borders5.top = xlwt.Borders.THIN    

    #无作用
    pattern = xlwt.Pattern()
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = 30  #boss的颜色
    
    #
    style_general = xlwt.XFStyle()
    style_general.font = font
    style_general.alignment = alignment

    #下框线，加粗，红色
    style_header = xlwt.XFStyle()
    style_header.alignment = alignment
    style_header.font = font_header
    style_header.borders = borders2

    #下框线，不加粗，黑色
    style_lastl = xlwt.XFStyle()
    style_lastl.alignment = alignment
    style_lastl.font = font
    style_lastl.borders = borders2

    #加粗，绿色，白字，下框线
    style_col = xlwt.XFStyle()
    style_col.font = font_col
    style_col.alignment = alignment
    style_col.borders = borders2
    style_col.pattern = pattern

    style_col_2 = xlwt.XFStyle()
    style_col_2.font = font_col
    style_col_2.alignment = alignment
    style_col_2.borders = borders4
    style_col_2.pattern = pattern

    #上框线 
    style_header_top = xlwt.XFStyle()
    style_header_top.borders = borders
    style_header_top.font = font_header
    style_header_top.alignment = alignment

    #下框线 
    style_index_b = xlwt.XFStyle()
    style_index_b.borders = borders4
    style_index_b.font = font_index
    style_index_b.alignment = alignment

    #上框线
    style_index_t = xlwt.XFStyle()
    style_index_t.borders = borders5
    style_index_t.font = font_index
    style_index_t.alignment = alignment

    #无b-t框线
    style_index_n = xlwt.XFStyle()
    style_index_n.borders = borders3
    style_index_n.font = font_index
    style_index_n.alignment = alignment

    style_r = xlwt.XFStyle()
    style_r.borders = borders3
    style_r.font = font
    style_r.alignment = alignment 

    return style_index_b, style_index_t, style_index_n, style_general, style_header, style_col, style_col_2, style_header_top, style_r, style_lastl



def setRules():
    #同比环比增幅，增量需要按照>=0，<0设置颜色规则，红色和绿色，全部字体bolt
    pass



#这个函数没用到，我是想直接从txt生成excel，不想通过pandas的df数据类型
def outputExcelByDF(excelName,sheetName,adf):
    workbook = xlwt.Workbook(encoding='utf-8')
    worksheet = workbook.add_sheet(sheetName)

    style_index_b, style_index_t, style_index_n, style_general, style_header, style_col, style_col_2, style_header_top, style_r, style_lastl = decorate()
    
    #在循环里面操作write这个函数
    nrows = adf.shape[0] 
    ncols = adf.shape[1]

    #逐个填入excel
    adfIndex = list(adf.index)

    for ind in range(len(adfIndex)):
        if adfIndex[ind] in ['ph1', 'ph2', 'ph3']:
            adfIndex[ind]  = ''
        if ind<11:
            if ind%3==1:
                worksheet.write(2+ind, 1, adfIndex[ind], style_index_b)
            else:
                worksheet.write(2+ind, 1, adfIndex[ind], style_index_t)
        elif ind == len(adfIndex)-1:
            worksheet.write(3+ind, 1, adfIndex[ind], style_index_b)
        elif ind>11:
            worksheet.write(3+ind, 1, adfIndex[ind], style_index_n)
         
    #worksheet.write_merge(1, 2, 0 ,0, '时间',style)
    for i in range(nrows):
        if i == 11:
            worksheet.write_merge(11+2,12+2,1,1,'时间',style_col_2) #蓝色，加粗，下框线
            worksheet.write_merge(11+2, 11+2, 2, 2+ncols//2, '规模', style_col_2) #蓝色，加粗，下框线
            worksheet.write_merge(11+2, 11+2, 3+ncols//2, 1+ncols, '占比', style_col) #蓝色，加粗，下框线
            #请使用txt文件传入12th的header
            for j in range(ncols):
                if j!= ncols//2:
                    worksheet.write(3+i, j+2, 'placeholder', style_col) #
                else:
                    worksheet.write(3+i, j+2, 'placeholder', style_col_2) #

        for j in range(ncols):
            cell_value = adf.iloc[i,j]

            #if isinstance(cell_value,int) is False and cell_value.isdigit() is True: 
            #    cell_value = int(cell_value)
            #elif cell_value == '': pass
            #
            #elif '%' in cell_value: 
            #    print(cell_value)
            #    cell_value = float(cell_value.replace('%',''))/100.0
            #elif 'N_' in cell_value: pass
            #else: cell_value = float(cell_value) 
            #worksheet.write(2+i, 1+j, cell_value,set_style('Times New Roman',220,True))

            if i<11: 
                if i%3 == 0: #需要border上框线
                    worksheet.write(2+i, j+2, cell_value, style_header_top) #
                elif i%3==1:
                    worksheet.write(2+i, j+2, cell_value, style_header) #
                else:
                    worksheet.write(2+i, j+2, cell_value, style_general) #无框线
            if i>11:
                if i == nrows-1 and j != ncols//2:
                    worksheet.write(3+i, j+2, cell_value, style_lastl) #下框线黑色字体
                elif i == nrows-1 and j == ncols//2:
                    worksheet.write(3+i, j+2, cell_value, style_index_b)
                elif j != ncols//2: 
                    worksheet.write(3+i, j+2, cell_value, style_general) #无框线
                elif j == ncols//2: 
                    worksheet.write(3+i, j+2, cell_value, style_r)

    workbook.save(excelName)



########################################################
#以下是修饰excel表所用，比如居中等                     
########################################################
def crossCheck():
    '''你应该把你计算ratio以及rate的结果用一种可靠的方式交叉检验一下
       因为在这个过程中真的很容易出错, 好了, 这个函数就是做这件事情的, 你也
       许需要几个函数或者类帮你完成这件事情, 这个事情真的很重要，它影响了
       你的分析师的reputation等各种东西。
    '''
    pass



#第二部分
#########################################################
#以下是用于从excel读取数据得到word文档所用, 
#这样就可以直接产生word了,不需要复制粘贴了
######################################################### 
def getIJ(rows, cols, sheet, acell):
    for i in range(rows):
        #if i !=5: continue
        for j in range(cols):
            cell = sheet.cell_value(i, j)
            if cell == acell: 
                #print(i+1, j+1)
                return i,j



def regularizeRate(aval):
    aval = float(aval)
    if aval >=0: 
        prefix = '+'
    else: 
        prefix = ''

    aval = str(round(aval*100,1))+'%'
    
    return prefix+aval
    #return aval



def regularizeRatio(aval):
    return regularizeRate(aval)   



def regularizeScale(aval):
    aval = int(aval)

    aval = abs(aval)

    if aval>=10000 and aval<100000:
        aval = str(round(aval/10000.0,1))+'W'
    elif aval<10000: 
       aval = str(aval)
    elif aval>=100000: 
        aval = str(round(aval/10000.0)) + 'W'

    #return prefix+aval
    return aval



def regularizeGrowth(aval):
    aval = int(aval)
    if aval >=0:
        prefix = '+'
    else:
        prefix = '-'

    aval = regularizeScale(aval)

    return prefix+aval



#########################################################################
#用于产生word文档的对象，以便自动生成word文件，而不需要
#########################################################################
def makeWordDoc(amonth):
    title = '活跃规模2019-{0}月报'.format(amonth)
    #wordName = '{}月报-活跃规模.docx'.format(amonth)
    document = Document()
    document.add_heading(title, level=1)

    #p = document.add_paragraph('A plain paragraph having some ')
    #document.add_paragraph('first item in unordered list', style='List Bullet')

    return document



def endWordDoc(adoc, amonth):
    wordName = '{}月报-活跃规模-update'.format(amonth)
    adoc.save('{0}.docx'.format(wordName))



def redOrGreen(it, aparag):
    run = aparag.add_run(it)
    font = run.font
    font.bold = True
    if '+' in it:
        font.color.rgb =  RGBColor(255, 0, 0)
    elif '-' in it:
        font.color.rgb = RGBColor(112, 173, 71)
    else:
        font.color.rgb = RGBColor(0, 0, 0)
    font.name = '苹果-简'
    return aparag



def blackAndThin(it, aparag):
    run = aparag.add_run(it)
    font = run.font
    font.name = '苹果-简'
    font.bold = False
    font.color.rgb = RGBColor(0, 0, 0)
    return aparag



def setStyle(alist,aparag):
    #'一线城市...'
    print(alist[0])
    aparag = blackAndThin(alist[0]+' ',aparag)
    aparag = redOrGreen(alist[1], aparag)

    #
    aparag = blackAndThin('，月环比',aparag)
    aparag = redOrGreen(alist[2], aparag)

    #
    aparag = blackAndThin('，增幅',aparag)
    aparag = redOrGreen(alist[3], aparag)

    #
    aparag = blackAndThin('，同期',aparag)
    aparag = redOrGreen(alist[4], aparag)

    #
    aparag = blackAndThin('，月环比',aparag)
    aparag = redOrGreen(alist[5], aparag)

    #
    aparag = blackAndThin('，增幅',aparag)
    aparag = redOrGreen(alist[6], aparag)

    #
    aparag = blackAndThin('；同比',aparag)
    aparag = redOrGreen(alist[7], aparag)

    aparag = blackAndThin('。',aparag)
    return aparag



#读取xlsx内容    
#########################################################################
# 读取excel表中的环比增幅，环比增量，同期环比增幅，同期环比增量，月同比 #
#########################################################################
def read_xlsx(worksheet, sheetname='Sheet1', adoc=None, atype=None):
   
    ##载入文件
    ##wb = openpyxl.load_workbook(filename)

    ##获取Sheet1工作表
    ##ws = wb.get_sheet_by_name(sheetname)
    ##sheet = worksheet.sheet_by_name(sheetname)

    ##按行读取
    ##for row in list(ws.rows)[2:]:
    ##    tmprow = row[1]
    ##    print(type(tmprow))
    ##    print(str(tmprow).split('.')[1])
    ##    if row != 18: continue
    ##    for cell in row[1:]:
    ##        print (cell.value)

    ##按列读
    ##for col in list(ws.columns)[1:]:
    ##    for cell in col[2:]:
    ##        print (cell.value)
    ##        print (type(cell))


    # 开始操作了啊，从这里开始
    #[XXX]多少，月环比[XXX]，增幅[XXX]，同期月环比[XXX]，增幅[XXX]，同比[XXX]
    #i, j，左上角的"环比增幅"字样的位置
    #列名的位置, arow, acol
    #sheet.cell_val(i,j)
    #print('{cell_value(arow, acol)}{cell_value(arow+1, acol)}, 月环比{cell_value(i,acol}, 增幅{cell_value(i+1, acol)}; 
    #      同期月环比{cell_value(i+3,acol)}, 增幅{cell_value(i+4, acol)}, 同比{cell_value(i+6, acol)} \n')

    sheet = worksheet.sheet_by_name(sheetname)

    rows = sheet.nrows
    cols = sheet.ncols

    iind, jind = getIJ(rows,cols,sheet,'环比增幅')
    arow, acol = getIJ(rows,cols,sheet,'时间')
    try:
        iend, jend = getIJ(rows,cols,sheet,'总量')
    except Exception:
        iend, jend = rows, cols


    arow = arow+1
    acol +=1

    #print(sheet.cell_value(arow,acol))
   
    #print(acol, jend)
    print(sheet.cell_value(iind-1,jind))
    subsubtitle = sheet.cell_value(iind-1,jind)

    sheetFeature = subsubtitle.split('-')[0]
    print(sheetFeature,'?==?',atype)
    if sheetFeature != atype:
        alevel = 3
        adoc.add_paragraph('')
        adoc.add_heading(sheetFeature, level=3)
        adoc.add_heading(subsubtitle, level=4)
    else: 
        adoc.add_heading(subsubtitle, level=4)
    atype = sheetFeature

    #run2.font.color.rgb = red

    for tmpj in range(acol, jend):
        acol = tmpj
        val1 = sheet.cell_value(arow, acol)
        val1 = str(val1)
        val2 = sheet.cell_value(arow+1, acol)
        val3 = sheet.cell_value(iind,acol)
        val4 = sheet.cell_value(iind+1, acol)

        val5 = sheet.cell_value(arow+13, acol)
        
        val6 = sheet.cell_value(iind+3,acol)
        val7 = sheet.cell_value(iind+4, acol)
        val8 = sheet.cell_value(iind+6, acol) 

        val2 = regularizeScale(val2)
        val3 = regularizeRate(val3)
        val4 = regularizeGrowth(val4)

        val5 = regularizeScale(val5)
        val6 = regularizeRate(val6)
        val7 = regularizeGrowth(val7)

        val8 = regularizeRate(val8)
     
      
        #write into MS doc 2        
        #print('{0} {1}，月环比{2}，增幅{3}；同期{4}，月环比{5}，增幅{6}；同比{7}；'.format(val1, val2, val3, val4, val5, val6, val7, val8) )
        #abullet = '{0} {1}，'.format(val1, val2)
        #月环比{2}，增幅{3}；同期{4}，月环比{5}，增幅{6}；同比{7}；'.format(val1, val2, val3, val4, val5, val6, val7, val8)
        #abullet = '{0} {1}，月环比{2}，增幅{3}；同期{4}，月环比{5}，增幅{6}；同比{7}；'.format(val1, val2, val3, val4, val5, val6, val7, val8)
        #adoc.add_paragraph(abullet, style='List Bullet')
        aparag = adoc.add_paragraph()
        #aparag = aparag.add_run()
        alist = list()

        alist.extend([val1, val2, val3, val4, val5, val6, val7, val8])
          
        aparag = setStyle(alist,aparag) 

    adoc.add_paragraph('')

    return atype
    #write into MS doc 3
    #print('\n\n')
    #adoc.add_paragraph('\n\n')
    #adoc.save('{0}.docx'.format(wordName))
    #之后的这里应该是调用wordWrite.py的脚本去得到WORD DOCUMENT
    #而不是在屏幕输出上command+c, command+v粘贴到word文档中
     


def formatOutputRes(excelName=None, adoc=None):
    filename = excelName
    #filename = '09月报-活跃规模.xlsx'
    filename = filename.strip()

    worksheet = xlrd.open_workbook(u'{}'.format(filename))
    sheet_names= worksheet.sheet_names()
    print(sheet_names,'\n\n')

    atype = ''
    #for sheetname in sheet_names:
    for sheetname in ['牛人-年龄','牛人-毕业院校']:
        print('*****************')
        print(sheetname)
        atype = read_xlsx(worksheet,sheetname, adoc, atype)
        print('*****************\n\n')



def parsePath(apath):
    apath = apath.split('/')
    sheetName = apath[-2]+'-'+apath[-1]
 
    if apath[-2] == 'geek' or 'boss' == apath[-2]:
        sheetName = 'active-completed-'+sheetName
    else:
        sheetName = 'active-online-'+sheetName
    
    return sheetName

 

#主函数
if __name__=='__main__':
    #告诉你的脚本，你想要计算那个月的月报数据
    #amonth=input('月报月份: ')

    #或许你可以使用datatime的当前时间的月份-1就是这里的amonth
    #print(amonth)
   
    #告诉这个脚本，你想要把哪个文件夹下面txt生成excel tables
    #我想这里很明显就是monthly/boss, 或者geek/salary之类的文件夹名称了
    #apath=sys.argv[1]

    #fs = os.listdir(apath)
    #for f in fs:
    #    if f.endswith('.txt') is False: continue
    #    excelName = f.split('geek/boss/expect/job/company')
    #    sheetName = f.split('city_level/l1_code/salary/...')

    #inputfilePath = sys.argv[1]
    
    #如果你要使用某个指定的txt文件，而不是某个文件夹下所有的txt文件的话
    #inputfileTxt = 'tmp.txt'

    #outfileExcel = '{}月报-{}.xlsx'.format(amonth,'活跃规模')
    #conver_txt_to_pd_to_xlsx(inputfileTxt,outfileExcel)

    #我注释了，先不产生excel
    #apath = './output/job/salary'
 
    #我注释了，先不产生excel
    #sheetName = parsePath(apath)

    #我注释了，先不产生excel
    #res = convert_txt_to_df(apath)
    ##print(type(res))
    ##print(sortDF(res))
    #我注释了，先不产生excel
    #res = sortDF(res)
     
    ##res = my_round(res)
    ##print(res)

    ##outputExcel('test-08月报-活跃规模.xlsx',sheetName,res)

    #我注释了，先不产生excel
    #outputExcelByDF('test-08月报-活跃规模.xls',sheetName,res)

    ##初始化 MS doc对象
    adoc = makeWordDoc('09')

    ##数据数据月报的描述段落 输出到MS doc中
    formatOutputRes('09月报-活跃规模.xlsx',adoc) 

    ##保存word文档
    endWordDoc(adoc,'09')



